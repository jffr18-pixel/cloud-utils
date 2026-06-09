"""Generacion del informe de revision del expediente.

Formatos disponibles:
  - generar_informe -> Markdown (cadena de texto)
  - generar_docx    -> Word (.docx) en bytes
  - generar_pdf     -> PDF en bytes
"""

import io
import os
from datetime import date

from . import tramites
from .analizador import calcular_permanencia, expediente_listo, incidencias_expediente

_EXT_LOGO_PDF = (".png", ".jpg", ".jpeg")
_EXT_LOGO_DOCX = (".png", ".jpg", ".jpeg", ".gif", ".bmp")


def _contacto(gestoria):
    """Linea de contacto a partir de los datos de la gestoria."""
    if not gestoria:
        return ""
    partes = [gestoria.get("direccion"), gestoria.get("telefono"), gestoria.get("email")]
    return " · ".join(p for p in partes if p)

_ICONO = {
    "correcto": "OK",
    "con_incidencias": "REVISAR",
    "proximo_a_caducar": "CADUCA PRONTO",
    "caducado": "CADUCADO",
    "falta": "FALTA",
    "falta_opcional": "FALTA (opcional)",
}

_ESTADO_DOC = {
    "vigente": "vigente",
    "caducado": "CADUCADO",
    "proximo_a_caducar": "proximo a caducar",
    "sin_caducidad": "sin caducidad",
    "ilegible": "ilegible",
    "desconocido": "sin determinar",
}


def generar_informe(checklist, no_identificados, tramite_id, solicitante="", hoy=None):
    """Devuelve el informe de revision como cadena Markdown."""
    if hoy is None:
        hoy = date.today()
    tramite = tramites.TRAMITES[tramite_id]
    listo = expediente_listo(checklist)

    lineas = []
    lineas.append(f"# Informe de revision de expediente")
    lineas.append("")
    lineas.append(f"- **Tramite:** {tramite['nombre']}")
    if solicitante:
        lineas.append(f"- **Solicitante:** {solicitante}")
    lineas.append(f"- **Fecha de revision:** {hoy.isoformat()}")
    veredicto = (
        "LISTO PARA PRESENTAR (revisar avisos)"
        if listo
        else "NO LISTO: faltan documentos obligatorios o hay caducados"
    )
    lineas.append(f"- **Resultado:** {veredicto}")
    lineas.append("")

    # Contadores
    faltan = [c for c in checklist if c["estado"] == "falta"]
    caducados = [c for c in checklist if c["estado"] == "caducado"]
    avisos = [c for c in checklist if c["estado"] in ("con_incidencias", "proximo_a_caducar")]

    lineas.append("## Resumen")
    lineas.append("")
    lineas.append(f"- Documentos obligatorios que faltan: **{len(faltan)}**")
    lineas.append(f"- Documentos caducados: **{len(caducados)}**")
    lineas.append(f"- Documentos a revisar / proximos a caducar: **{len(avisos)}**")
    lineas.append("")

    # Comprobaciones automaticas
    lineas.append("## Comprobaciones automaticas")
    lineas.append("")
    for linea in _lineas_comprobaciones(checklist, no_identificados, tramite_id, hoy):
        lineas.append(f"- {linea}")
    lineas.append("")

    # Checklist
    lineas.append("## Checklist de documentacion")
    lineas.append("")
    lineas.append("| Estado | Documento | Obligatorio | Detalle |")
    lineas.append("|---|---|---|---|")
    for fila in checklist:
        icono = _ICONO.get(fila["estado"], fila["estado"])
        oblig = "Si" if fila["obligatorio"] else "No"
        detalle = _detalle_fila(fila)
        lineas.append(f"| {icono} | {fila['nombre']} | {oblig} | {detalle} |")
    lineas.append("")

    # Detalle por documento aportado
    lineas.append("## Detalle de los documentos aportados")
    lineas.append("")
    aportados = [d for fila in checklist for d in fila["documentos"]] + no_identificados
    if not aportados:
        lineas.append("_No se ha aportado ningun documento._")
        lineas.append("")
    for doc in aportados:
        lineas.extend(_bloque_documento(doc))

    # Documentos no identificados
    if no_identificados:
        lineas.append("## Documentos no identificados")
        lineas.append("")
        lineas.append(
            "Estos archivos no encajan con ningun documento exigido por el tramite. "
            "Revisar manualmente:"
        )
        for doc in no_identificados:
            lineas.append(f"- {doc.get('archivo', '-')}: {doc.get('resumen', '')}")
        lineas.append("")

    # Acciones recomendadas
    lineas.append("## Acciones recomendadas")
    lineas.append("")
    acciones = _acciones(faltan, caducados, avisos)
    if not acciones:
        lineas.append("- Sin acciones pendientes. Expediente completo.")
    else:
        lineas.extend(f"- {a}" for a in acciones)
    lineas.append("")

    lineas.append("---")
    lineas.append(
        "_Informe generado automaticamente como apoyo a la revision. "
        "No sustituye el criterio profesional del gestor._"
    )
    return "\n".join(lineas)


def _detalle_fila(fila):
    if fila["estado"] in ("falta", "falta_opcional"):
        return fila.get("notas", "") or "No aportado."
    partes = []
    for doc in fila["documentos"]:
        info = doc.get("archivo", "documento")
        if doc.get("fecha_caducidad"):
            info += f" (cad. {doc['fecha_caducidad']})"
        partes.append(info)
    incidencias = [i for doc in fila["documentos"] for i in doc.get("incidencias", [])]
    detalle = "; ".join(partes)
    if incidencias:
        detalle += " — " + "; ".join(incidencias)
    return detalle or "Aportado."


def _bloque_documento(doc):
    lineas = []
    titulo = doc.get("tipo_nombre") or doc.get("archivo", "Documento")
    lineas.append(f"### {titulo}")
    lineas.append(f"- Archivo: {doc.get('archivo', '-')}")
    if doc.get("titular"):
        lineas.append(f"- Titular: {doc['titular']}")
    if doc.get("numero"):
        lineas.append(f"- Numero: {doc['numero']}")
    if doc.get("pais_emision"):
        lineas.append(f"- Pais/autoridad: {doc['pais_emision']}")
    if doc.get("fecha_emision"):
        lineas.append(f"- Fecha de emision: {doc['fecha_emision']}")
    if doc.get("fecha_caducidad"):
        lineas.append(f"- Fecha de caducidad: {doc['fecha_caducidad']}")
    lineas.append(f"- Estado: {_ESTADO_DOC.get(doc.get('estado'), doc.get('estado', '-'))}")
    lineas.append(f"- Legibilidad: {doc.get('legibilidad', '-')}")
    if doc.get("incidencias"):
        lineas.append("- Incidencias:")
        lineas.extend(f"    - {i}" for i in doc["incidencias"])
    if doc.get("resumen"):
        lineas.append(f"- Resumen: {doc['resumen']}")
    lineas.append("")
    return lineas


def _acciones(faltan, caducados, avisos):
    acciones = []
    for fila in faltan:
        acciones.append(f"Solicitar al cliente: {fila['nombre']}.")
    for fila in caducados:
        acciones.append(f"Renovar (documento caducado): {fila['nombre']}.")
    for fila in avisos:
        for doc in fila["documentos"]:
            for inc in doc.get("incidencias", []):
                acciones.append(f"{fila['nombre']}: {inc}")
    # Eliminar duplicados conservando orden
    vistas = set()
    unicas = []
    for a in acciones:
        if a not in vistas:
            vistas.add(a)
            unicas.append(a)
    return unicas


def _contexto(checklist, tramite_id, solicitante, hoy):
    """Datos comunes a todos los formatos del informe."""
    if hoy is None:
        hoy = date.today()
    return {
        "tramite": tramites.TRAMITES[tramite_id]["nombre"],
        "solicitante": solicitante,
        "hoy": hoy,
        "listo": expediente_listo(checklist),
        "faltan": [c for c in checklist if c["estado"] == "falta"],
        "caducados": [c for c in checklist if c["estado"] == "caducado"],
        "avisos": [
            c for c in checklist if c["estado"] in ("con_incidencias", "proximo_a_caducar")
        ],
    }


def _todos_documentos(checklist, no_identificados):
    return [d for fila in checklist for d in fila["documentos"]] + no_identificados


def _lineas_comprobaciones(checklist, no_identificados, tramite_id, hoy):
    """Lineas de texto con coherencia y permanencia (comunes a los formatos)."""
    docs = _todos_documentos(checklist, no_identificados)
    lineas = []

    incidencias = incidencias_expediente(docs, hoy)
    if incidencias:
        for aviso in incidencias:
            lineas.append(f"Aviso: {aviso}")
    else:
        lineas.append("Coherencia: nombres, numeros y fechas coinciden entre los documentos.")

    permanencia = calcular_permanencia(docs, tramite_id, hoy=hoy)
    if permanencia:
        req = permanencia["requeridos"]
        if permanencia["anios"] is None:
            lineas.append(
                f"Permanencia: este tramite exige {req} anos; no se ha podido determinar "
                "la fecha de inicio a partir de los documentos."
            )
        else:
            estado = "CUMPLE" if permanencia["cumple"] else "NO CUMPLE"
            lineas.append(
                f"Permanencia: {permanencia['anios']} anos acreditados desde "
                f"{permanencia['fecha_inicio']} (exigidos {req}) -> {estado}."
            )
    return lineas


# --------------------------------------------------------------------------- #
#  Word (.docx)
# --------------------------------------------------------------------------- #
def generar_docx(
    checklist, no_identificados, tramite_id, solicitante="", hoy=None, gestoria=None
):
    """Genera el informe en formato Word y lo devuelve como bytes."""
    from docx import Document  # import diferido para no exigir la dependencia siempre
    from docx.shared import Inches, Pt

    ctx = _contexto(checklist, tramite_id, solicitante, hoy)
    doc = Document()

    # Membrete de la gestoria
    if gestoria:
        logo = gestoria.get("logo_path", "")
        if logo and os.path.exists(logo) and logo.lower().endswith(_EXT_LOGO_DOCX):
            try:
                doc.add_picture(logo, width=Inches(1.6))
            except Exception:  # noqa: BLE001
                pass
        if gestoria.get("nombre_gestoria"):
            p = doc.add_paragraph()
            run = p.add_run(gestoria["nombre_gestoria"])
            run.bold = True
            run.font.size = Pt(14)
        contacto = _contacto(gestoria)
        if contacto:
            doc.add_paragraph(contacto)
        if logo or gestoria.get("nombre_gestoria") or contacto:
            doc.add_paragraph("")

    doc.add_heading("Informe de revision de expediente", level=0)
    doc.add_paragraph(f"Tramite: {ctx['tramite']}")
    if solicitante:
        doc.add_paragraph(f"Solicitante: {solicitante}")
    doc.add_paragraph(f"Fecha de revision: {ctx['hoy'].isoformat()}")
    veredicto = (
        "LISTO PARA PRESENTAR (revisar avisos)"
        if ctx["listo"]
        else "NO LISTO: faltan documentos obligatorios o hay caducados"
    )
    p = doc.add_paragraph()
    p.add_run("Resultado: ").bold = True
    p.add_run(veredicto)

    doc.add_heading("Resumen", level=1)
    doc.add_paragraph(f"Documentos obligatorios que faltan: {len(ctx['faltan'])}")
    doc.add_paragraph(f"Documentos caducados: {len(ctx['caducados'])}")
    doc.add_paragraph(f"Documentos a revisar / proximos a caducar: {len(ctx['avisos'])}")

    doc.add_heading("Comprobaciones automaticas", level=1)
    for linea in _lineas_comprobaciones(checklist, no_identificados, tramite_id, ctx["hoy"]):
        doc.add_paragraph(linea, style="List Bullet")

    doc.add_heading("Checklist de documentacion", level=1)
    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    cab = tabla.rows[0].cells
    cab[0].text, cab[1].text, cab[2].text, cab[3].text = (
        "Estado",
        "Documento",
        "Obligatorio",
        "Detalle",
    )
    for fila in checklist:
        celdas = tabla.add_row().cells
        celdas[0].text = _ICONO.get(fila["estado"], fila["estado"])
        celdas[1].text = fila["nombre"]
        celdas[2].text = "Si" if fila["obligatorio"] else "No"
        celdas[3].text = _detalle_fila(fila)

    doc.add_heading("Detalle de los documentos aportados", level=1)
    aportados = [d for fila in checklist for d in fila["documentos"]] + no_identificados
    if not aportados:
        doc.add_paragraph("No se ha aportado ningun documento.")
    for docu in aportados:
        doc.add_heading(docu.get("tipo_nombre") or docu.get("archivo", "Documento"), level=2)
        for linea in _lineas_documento_plano(docu):
            doc.add_paragraph(linea, style="List Bullet")

    if no_identificados:
        doc.add_heading("Documentos no identificados", level=1)
        for docu in no_identificados:
            doc.add_paragraph(
                f"{docu.get('archivo', '-')}: {docu.get('resumen', '')}", style="List Bullet"
            )

    doc.add_heading("Acciones recomendadas", level=1)
    acciones = _acciones(ctx["faltan"], ctx["caducados"], ctx["avisos"])
    if not acciones:
        doc.add_paragraph("Sin acciones pendientes. Expediente completo.")
    for accion in acciones:
        doc.add_paragraph(accion, style="List Bullet")

    doc.add_paragraph(
        "Informe generado automaticamente como apoyo a la revision. "
        "No sustituye el criterio profesional del gestor."
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
#  PDF
# --------------------------------------------------------------------------- #
def _latin1(texto):
    """Adapta el texto a la codificacion de las fuentes base del PDF."""
    if not texto:
        return ""
    reemplazos = {
        "—": "-",
        "–": "-",
        "‘": "'",
        "’": "'",
        "“": '"',
        "”": '"',
        "…": "...",
        " ": " ",
        "€": "EUR",
    }
    for origen, destino in reemplazos.items():
        texto = texto.replace(origen, destino)
    return texto.encode("latin-1", "replace").decode("latin-1")


def _cabecera_pdf(pdf, gestoria):
    """Dibuja el membrete de la gestoria en la parte superior del PDF."""
    if not gestoria:
        return
    logo = gestoria.get("logo_path", "")
    tiene_logo = bool(logo and os.path.exists(logo) and logo.lower().endswith(_EXT_LOGO_PDF))
    nombre = gestoria.get("nombre_gestoria", "")
    contacto = _contacto(gestoria)
    if not (tiene_logo or nombre or contacto):
        return
    if tiene_logo:
        try:
            pdf.image(logo, x=10, y=8, w=24)
        except Exception:  # noqa: BLE001
            tiene_logo = False
    x_text = 38 if tiene_logo else 10
    pdf.set_xy(x_text, 10)
    if nombre:
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_x(x_text)
        pdf.multi_cell(0, 6, _latin1(nombre), new_x="LMARGIN", new_y="NEXT")
    if contacto:
        pdf.set_x(x_text)
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(0, 5, _latin1(contacto), new_x="LMARGIN", new_y="NEXT")
    y = max(pdf.get_y(), 30)
    pdf.set_draw_color(160, 160, 160)
    pdf.line(10, y, 200, y)
    pdf.set_y(y + 4)


def generar_pdf(
    checklist, no_identificados, tramite_id, solicitante="", hoy=None, gestoria=None
):
    """Genera el informe en formato PDF y lo devuelve como bytes."""
    from fpdf import FPDF  # import diferido

    ctx = _contexto(checklist, tramite_id, solicitante, hoy)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _cabecera_pdf(pdf, gestoria)

    def titulo(texto, tam=14):
        pdf.set_font("Helvetica", "B", tam)
        pdf.multi_cell(0, 7, _latin1(texto), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)

    def parrafo(texto, negrita=False):
        pdf.set_font("Helvetica", "B" if negrita else "", 10)
        pdf.multi_cell(0, 5, _latin1(texto), new_x="LMARGIN", new_y="NEXT")

    titulo("Informe de revision de expediente", 16)
    parrafo(f"Tramite: {ctx['tramite']}")
    if solicitante:
        parrafo(f"Solicitante: {solicitante}")
    parrafo(f"Fecha de revision: {ctx['hoy'].isoformat()}")
    veredicto = (
        "LISTO PARA PRESENTAR (revisar avisos)"
        if ctx["listo"]
        else "NO LISTO: faltan documentos obligatorios o hay caducados"
    )
    parrafo(f"Resultado: {veredicto}", negrita=True)
    pdf.ln(2)

    titulo("Resumen")
    parrafo(f"- Documentos obligatorios que faltan: {len(ctx['faltan'])}")
    parrafo(f"- Documentos caducados: {len(ctx['caducados'])}")
    parrafo(f"- A revisar / proximos a caducar: {len(ctx['avisos'])}")
    pdf.ln(2)

    titulo("Comprobaciones automaticas")
    for linea in _lineas_comprobaciones(checklist, no_identificados, tramite_id, ctx["hoy"]):
        parrafo(f"- {linea}")
    pdf.ln(2)

    titulo("Checklist de documentacion")
    pdf.set_font("Helvetica", "", 9)
    with pdf.table(col_widths=(16, 44, 14, 46), text_align="LEFT") as tabla:
        encabezado = tabla.row()
        for celda in ("Estado", "Documento", "Oblig.", "Detalle"):
            encabezado.cell(celda)
        for fila in checklist:
            row = tabla.row()
            row.cell(_latin1(_ICONO.get(fila["estado"], fila["estado"])))
            row.cell(_latin1(fila["nombre"]))
            row.cell("Si" if fila["obligatorio"] else "No")
            row.cell(_latin1(_detalle_fila(fila)))
    pdf.ln(2)

    titulo("Detalle de los documentos aportados")
    aportados = [d for fila in checklist for d in fila["documentos"]] + no_identificados
    if not aportados:
        parrafo("No se ha aportado ningun documento.")
    for docu in aportados:
        parrafo(docu.get("tipo_nombre") or docu.get("archivo", "Documento"), negrita=True)
        for linea in _lineas_documento_plano(docu):
            parrafo(f"  - {linea}")
        pdf.ln(1)

    if no_identificados:
        titulo("Documentos no identificados")
        for docu in no_identificados:
            parrafo(f"- {docu.get('archivo', '-')}: {docu.get('resumen', '')}")
        pdf.ln(2)

    titulo("Acciones recomendadas")
    acciones = _acciones(ctx["faltan"], ctx["caducados"], ctx["avisos"])
    if not acciones:
        parrafo("- Sin acciones pendientes. Expediente completo.")
    for accion in acciones:
        parrafo(f"- {accion}")
    pdf.ln(3)

    pdf.set_font("Helvetica", "I", 8)
    pdf.multi_cell(
        0,
        4,
        _latin1(
            "Informe generado automaticamente como apoyo a la revision. "
            "No sustituye el criterio profesional del gestor."
        ),
        new_x="LMARGIN",
        new_y="NEXT",
    )

    salida = pdf.output()
    return bytes(salida)


def _pendientes(checklist):
    """Devuelve (faltan, caducados, proximos) como listas de nombres/textos."""
    faltan = [c["nombre"] for c in checklist if c["estado"] == "falta"]
    caducados = []
    proximos = []
    for c in checklist:
        if c["estado"] == "caducado":
            caducados.append(c["nombre"])
        elif c["estado"] == "proximo_a_caducar":
            for d in c["documentos"]:
                cad = d.get("fecha_caducidad")
                proximos.append(f"{c['nombre']} (caduca el {cad})" if cad else c["nombre"])
    return faltan, caducados, proximos


def _texto_requerimiento(checklist, tramite_id, solicitante, gestoria, hoy):
    """Cuerpo de la carta de requerimiento como lista de parrafos (texto plano)."""
    if hoy is None:
        hoy = date.today()
    tramite = tramites.TRAMITES[tramite_id]["nombre"]
    faltan, caducados, proximos = _pendientes(checklist)
    saludo = f"Estimado/a {solicitante}:" if solicitante else "Estimado/a cliente:"

    parrafos = [
        f"Fecha: {hoy.isoformat()}",
        "",
        saludo,
        "",
        f"En relacion con su tramite de {tramite}, tras revisar la documentacion "
        "aportada le informamos de lo siguiente:",
    ]
    if not (faltan or caducados or proximos):
        parrafos += [
            "",
            "La documentacion esta completa. No es necesario que aporte nada mas por el momento.",
        ]
    if faltan:
        parrafos += ["", "Documentacion pendiente de aportar:"]
        parrafos += [f"  - {n}" for n in faltan]
    if caducados:
        parrafos += ["", "Documentacion caducada que debe renovar y volver a aportar:"]
        parrafos += [f"  - {n}" for n in caducados]
    if proximos:
        parrafos += ["", "Documentacion proxima a caducar (conviene renovarla cuanto antes):"]
        parrafos += [f"  - {n}" for n in proximos]

    parrafos += [
        "",
        "Le rogamos nos haga llegar la documentacion indicada a la mayor brevedad "
        "para poder continuar con la tramitacion de su expediente.",
        "",
        "Quedamos a su disposicion para cualquier aclaracion.",
        "",
        "Atentamente,",
    ]
    if gestoria and gestoria.get("nombre_gestoria"):
        parrafos.append(gestoria["nombre_gestoria"])
    contacto = _contacto(gestoria)
    if contacto:
        parrafos.append(contacto)
    return parrafos


def generar_requerimiento(checklist, tramite_id, solicitante="", gestoria=None, hoy=None):
    """Carta de requerimiento al cliente en texto plano (lista de lo pendiente)."""
    return "\n".join(_texto_requerimiento(checklist, tramite_id, solicitante, gestoria, hoy))


def generar_requerimiento_docx(checklist, tramite_id, solicitante="", gestoria=None,
                               hoy=None, firma_imagen=None):
    """Carta de requerimiento en formato Word (con membrete si esta configurado).

    Si se aporta 'firma_imagen' (bytes de una imagen PNG/JPG), se incorpora al
    final del documento como constancia de conformidad firmada por el cliente.
    """
    from docx import Document
    from docx.shared import Inches, Pt

    doc = Document()
    if gestoria:
        logo = gestoria.get("logo_path", "")
        if logo and os.path.exists(logo) and logo.lower().endswith(_EXT_LOGO_DOCX):
            try:
                doc.add_picture(logo, width=Inches(1.6))
            except Exception:  # noqa: BLE001
                pass
        if gestoria.get("nombre_gestoria"):
            p = doc.add_paragraph()
            run = p.add_run(gestoria["nombre_gestoria"])
            run.bold = True
            run.font.size = Pt(14)
        doc.add_paragraph("")

    for parrafo in _texto_requerimiento(checklist, tramite_id, solicitante, gestoria, hoy):
        doc.add_paragraph(parrafo)

    if firma_imagen:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Firma del cliente (conforme):")
        run.bold = True
        try:
            doc.add_picture(io.BytesIO(firma_imagen), width=Inches(2.2))
        except Exception:  # noqa: BLE001
            pass
        if solicitante:
            doc.add_paragraph(solicitante)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _lineas_documento_plano(doc):
    """Lineas de detalle de un documento, en texto plano (para docx y pdf)."""
    lineas = [f"Archivo: {doc.get('archivo', '-')}"]
    if doc.get("titular"):
        lineas.append(f"Titular: {doc['titular']}")
    if doc.get("numero"):
        lineas.append(f"Numero: {doc['numero']}")
    if doc.get("pais_emision"):
        lineas.append(f"Pais/autoridad: {doc['pais_emision']}")
    if doc.get("fecha_emision"):
        lineas.append(f"Fecha de emision: {doc['fecha_emision']}")
    if doc.get("fecha_caducidad"):
        lineas.append(f"Fecha de caducidad: {doc['fecha_caducidad']}")
    lineas.append(f"Estado: {_ESTADO_DOC.get(doc.get('estado'), doc.get('estado', '-'))}")
    lineas.append(f"Legibilidad: {doc.get('legibilidad', '-')}")
    for inc in doc.get("incidencias", []):
        lineas.append(f"Incidencia: {inc}")
    if doc.get("resumen"):
        lineas.append(f"Resumen: {doc['resumen']}")
    return lineas


# --------------------------------------------------------------------------- #
#  Lista de documentacion exigida (para entregar al cliente)
# --------------------------------------------------------------------------- #
def generar_lista_docs_docx(tramite_id, solicitante="", gestoria=None, estado_docs=None):
    """Word con la lista de documentos exigidos para el tramite.

    ``estado_docs`` es un dict {tipo_id: estado} con el estado de cada doc
    en el expediente; si se pasa, se muestra un tick/cruz junto a cada linea.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # Margenes mas pequeños
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cabecera gestoria
    if gestoria:
        logo = gestoria.get("logo_path", "")
        if logo and os.path.exists(logo) and logo.lower().endswith(_EXT_LOGO_DOCX):
            try:
                doc.add_picture(logo, width=Inches(1.4))
            except Exception:
                pass
        nombre_g = gestoria.get("nombre_gestoria", "")
        contacto = _contacto(gestoria)
        if nombre_g:
            p = doc.add_paragraph(nombre_g)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(11)
        if contacto:
            doc.add_paragraph(contacto).runs[0].font.size = Pt(9)
        doc.add_paragraph("")

    tramite_datos = tramites.TRAMITES.get(tramite_id, {})
    tramite_nombre = tramite_datos.get("nombre", tramite_id)

    titulo = doc.add_heading(f"Documentacion necesaria: {tramite_nombre}", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if solicitante:
        p = doc.add_paragraph(f"Cliente: {solicitante}")
        p.runs[0].font.size = Pt(10)
    p = doc.add_paragraph(f"Fecha: {date.today().strftime('%d/%m/%Y')}")
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph("")

    docs_obligatorios = [d for d in tramites.documentos_de(tramite_id) if d["obligatorio"]]
    docs_opcionales   = [d for d in tramites.documentos_de(tramite_id) if not d["obligatorio"]]

    estado_docs = estado_docs or {}

    def _marca(tipo_id):
        est = estado_docs.get(tipo_id)
        if est == "correcto":
            return "✅ "
        if est in ("caducado", "falta"):
            return "❌ "
        if est in ("con_incidencias", "proximo_a_caducar"):
            return "⚠️ "
        return "☐ "

    doc.add_heading("Documentos obligatorios", level=2)
    for d in docs_obligatorios:
        marca = _marca(d["id"])
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{marca}{d['nombre']}")
        run.bold = True
        if d.get("notas"):
            p.add_run(f"\n   {d['notas']}").font.size = Pt(9)

    if docs_opcionales:
        doc.add_paragraph("")
        doc.add_heading("Documentos opcionales / complementarios", level=2)
        for d in docs_opcionales:
            marca = _marca(d["id"])
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{marca}{d['nombre']}")
            if d.get("notas"):
                p.add_run(f"\n   {d['notas']}").font.size = Pt(9)

    doc.add_paragraph("")
    nota = doc.add_paragraph(
        "Nota: los documentos originales deben estar en vigor en la fecha de presentacion. "
        "Entregue siempre copia junto al original para su cotejo."
    )
    nota.runs[0].font.size = Pt(9)
    nota.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
#  Ficha completa del cliente — PDF de una pagina (resumen imprimible)
# --------------------------------------------------------------------------- #
def generar_ficha_completa_pdf(registro, gestoria=None):
    """PDF de una pagina con los datos del cliente, estado del expediente y tareas."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _cabecera_pdf(pdf, gestoria)

    def tit(t):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_fill_color(230, 220, 250)
        pdf.multi_cell(0, 6, _latin1(t), fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(1)

    def lin(etiq, val):
        if not val:
            return
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(48, 5, _latin1(etiq + ":"))
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(0, 5, _latin1(str(val)), new_x="LMARGIN", new_y="NEXT")

    tramite_n = tramites.TRAMITES.get(registro.get("tramite_id", ""), {}).get(
        "nombre", registro.get("tramite_id", "-")
    )
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(0, 7, _latin1(f"Ficha: {registro.get('solicitante') or '-'}"),
                   new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(
        0, 5,
        _latin1(f"Tramite: {tramite_n}  |  Revision: {registro.get('fecha','')[:10]}"),
        new_x="LMARGIN", new_y="NEXT",
    )
    pdf.ln(3)

    tit("Datos personales")
    lin("Nombre", registro.get("nombre") or registro.get("solicitante"))
    lin("F. nacimiento", registro.get("fecha_nacimiento"))
    lin("Nacionalidad", registro.get("nacionalidad"))
    lin("NIE", registro.get("nie"))
    lin("Pasaporte", registro.get("num_pasaporte"))
    lin("Cad. pasaporte", registro.get("cad_pasaporte"))
    lin("Entrada en Espana", registro.get("fecha_entrada_espana"))
    lin("Telefono", registro.get("telefono"))
    lin("Email", registro.get("email_cliente"))
    lin("Direccion", f"{registro.get('direccion','') or ''} {registro.get('ciudad','') or ''}".strip())
    lin("Empleador", registro.get("empleador"))
    lin("Contrato", f"{registro.get('tipo_contrato','') or ''} {registro.get('fecha_contrato','') or ''}".strip())
    pdf.ln(2)

    tit("Estado del expediente")
    lin("Nº expediente", registro.get("numero_expediente"))
    lin("Resultado", registro.get("resultado_final", "pendiente"))
    lin("Presentado", "Si" if registro.get("presentado") else "No")
    h = registro.get("honorarios") or {}
    if h.get("importe", 0):
        lin("Honorarios", f"{h.get('importe',0):.2f} EUR  (cobrado: {h.get('cobrado',0):.2f} EUR)")
    pdf.ln(2)

    try:
        from .analizador import evaluar_expediente
        checklist, _ = evaluar_expediente(
            registro.get("resultados", []), registro.get("tramite_id", "")
        )
    except Exception:
        checklist = []

    if checklist:
        tit("Documentacion")
        _IC = {"correcto": "OK", "falta": "FALTA", "caducado": "CADUCADO",
               "proximo_a_caducar": "PRONTO", "con_incidencias": "REVISAR",
               "falta_opcional": "OPCIONAL"}
        pdf.set_font("Helvetica", "", 9)
        for item in checklist:
            icono = _IC.get(item["estado"], item["estado"])
            pdf.cell(22, 5, _latin1(f"[{icono}]"))
            pdf.multi_cell(0, 5, _latin1(item["nombre"]), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    seg = registro.get("seguimiento", [])
    if seg:
        tit("Seguimiento")
        for ev in seg[-5:]:
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(24, 5, _latin1(ev.get("fecha", "")[:10]))
            pdf.set_font("Helvetica", "", 9)
            pdf.multi_cell(0, 5, _latin1(f"{ev.get('estado','')} {ev.get('nota','')}".strip()),
                           new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    tareas_pend = [t for t in registro.get("tareas", []) if not t.get("hecha")]
    if tareas_pend:
        tit("Tareas pendientes")
        pdf.set_font("Helvetica", "", 9)
        for t in tareas_pend[:6]:
            pdf.multi_cell(0, 5, _latin1(f"  [ ] {t.get('fecha','')} — {t.get('descripcion','')}"),
                           new_x="LMARGIN", new_y="NEXT")

    if registro.get("notas"):
        pdf.ln(2)
        tit("Notas")
        pdf.set_font("Helvetica", "", 9)
        pdf.multi_cell(0, 5, _latin1(registro["notas"]), new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", "I", 8)
    pdf.set_y(-18)
    pdf.multi_cell(0, 4, _latin1("Documento generado automaticamente — uso interno de la gestoria."),
                   new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())


# --------------------------------------------------------------------------- #
#  Hoja de encargo / contrato de servicios — Word
# --------------------------------------------------------------------------- #
def generar_hoja_encargo_docx(registro, gestoria=None):
    """Hoja de encargo profesional pre-rellenada con datos del cliente y gestoria."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    h = registro.get("honorarios") or {}
    importe = float(h.get("importe", 0))
    concepto = h.get("concepto", "")
    solicitante = registro.get("nombre") or registro.get("solicitante") or "el/la cliente"
    tramite_n = tramites.TRAMITES.get(registro.get("tramite_id", ""), {}).get(
        "nombre", registro.get("tramite_id", "-")
    )
    hoy_str = date.today().strftime("%d/%m/%Y")
    nombre_g = (gestoria or {}).get("nombre_gestoria", "la gestoria")
    contacto_g = _contacto(gestoria or {})

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    if gestoria:
        logo = gestoria.get("logo_path", "")
        if logo and os.path.exists(logo) and logo.lower().endswith(_EXT_LOGO_DOCX):
            try:
                doc.add_picture(logo, width=Inches(1.4))
            except Exception:
                pass
        if gestoria.get("nombre_gestoria"):
            p = doc.add_paragraph(gestoria["nombre_gestoria"])
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(13)
        if contacto_g:
            doc.add_paragraph(contacto_g).runs[0].font.size = Pt(9)
        doc.add_paragraph("")

    h1 = doc.add_heading("HOJA DE ENCARGO PROFESIONAL", level=1)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha = doc.add_paragraph(f"Fecha: {hoy_str}")
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph("")
    doc.add_paragraph(
        f"De una parte, {nombre_g}" +
        (f" ({contacto_g})" if contacto_g else "") +
        ", en adelante EL PROFESIONAL."
    )
    doc.add_paragraph(f"De otra parte, D./Dna. {solicitante}, en adelante EL CLIENTE.")
    nie = registro.get("nie", "")
    pas = registro.get("num_pasaporte", "")
    id_doc = f"NIE: {nie}" if nie else (f"Pasaporte: {pas}" if pas else "")
    if id_doc:
        doc.add_paragraph(f"Documento de identidad: {id_doc}")
    if registro.get("telefono"):
        doc.add_paragraph(f"Telefono de contacto: {registro['telefono']}")
    if registro.get("email_cliente"):
        doc.add_paragraph(f"Email: {registro['email_cliente']}")

    doc.add_paragraph("")
    doc.add_heading("OBJETO DEL ENCARGO", level=2)
    doc.add_paragraph(
        f"El cliente encarga a {nombre_g} la tramitacion del procedimiento de "
        f"{tramite_n} ante las autoridades competentes, incluyendo la preparacion "
        "de la documentacion, revision del expediente, presentacion y seguimiento."
    )
    if concepto:
        doc.add_paragraph(f"Descripcion adicional: {concepto}")

    doc.add_paragraph("")
    doc.add_heading("HONORARIOS", level=2)
    if importe:
        doc.add_paragraph(
            f"Los honorarios acordados ascienden a {importe:.2f} EUR "
            "(IVA no incluido salvo indicacion contraria)."
        )
    else:
        doc.add_paragraph("Los honorarios se acordaran por separado.")

    doc.add_paragraph("")
    doc.add_heading("OBLIGACIONES DEL CLIENTE", level=2)
    for ob in [
        "Aportar toda la documentacion necesaria en tiempo y forma.",
        "Comunicar cualquier cambio de domicilio, situacion laboral o datos personales.",
        "Conservar copias de todos los documentos entregados.",
    ]:
        doc.add_paragraph(ob, style="List Bullet")

    doc.add_paragraph("")
    doc.add_heading("PROTECCION DE DATOS", level=2)
    doc.add_paragraph(
        f"Los datos personales facilitados seran tratados por {nombre_g} conforme al "
        "RGPD (Reglamento UE 2016/679) y la LOPDGDD."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")
    sig = doc.add_table(rows=1, cols=2)
    sig.cell(0, 0).text = f"Firma del profesional\n\n\n\n{nombre_g}"
    sig.cell(0, 1).text = f"Firma del cliente\n\n\n\n{solicitante}"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
#  Presupuesto / nota de honorarios — PDF
# --------------------------------------------------------------------------- #
def generar_presupuesto_pdf(registro, gestoria=None):
    """PDF de presupuesto/factura con membrete, datos del cliente y honorarios."""
    from fpdf import FPDF

    h = registro.get("honorarios") or {}
    importe = float(h.get("importe", 0))
    cobrado = float(h.get("cobrado", 0))
    pendiente = round(importe - cobrado, 2)
    tramite_n = tramites.TRAMITES.get(registro.get("tramite_id", ""), {}).get("nombre", "Tramite")
    concepto = h.get("concepto", "") or tramite_n
    solicitante = registro.get("nombre") or registro.get("solicitante") or "Cliente"
    hoy_str = date.today().strftime("%d/%m/%Y")
    ref = (registro.get("numero_expediente") or registro.get("id", ""))[:16]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _cabecera_pdf(pdf, gestoria)

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 8, "PRESUPUESTO / NOTA DE HONORARIOS", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 5, _latin1(f"Referencia: {ref}  |  Fecha: {hoy_str}"),
                   new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(0, 6, "CLIENTE", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 5, _latin1(solicitante), new_x="LMARGIN", new_y="NEXT")
    nie = registro.get("nie", "")
    pas = registro.get("num_pasaporte", "")
    id_doc = f"NIE: {nie}" if nie else (f"Pasaporte: {pas}" if pas else "")
    if id_doc:
        pdf.multi_cell(0, 5, _latin1(id_doc), new_x="LMARGIN", new_y="NEXT")
    if registro.get("email_cliente"):
        pdf.multi_cell(0, 5, _latin1(registro["email_cliente"]), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(0, 6, "SERVICIOS", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    with pdf.table(col_widths=(100, 30, 30, 30), text_align="LEFT") as tabla:
        enc = tabla.row()
        for hdr in ("Concepto", "Importe", "Cobrado", "Pendiente"):
            enc.cell(hdr)
        row = tabla.row()
        row.cell(_latin1(concepto))
        row.cell(f"{importe:.2f} EUR")
        row.cell(f"{cobrado:.2f} EUR")
        row.cell(f"{pendiente:.2f} EUR")
    pdf.ln(4)

    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(0, 6, _latin1(f"TOTAL: {importe:.2f} EUR"), new_x="LMARGIN", new_y="NEXT")
    if pendiente > 0:
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _latin1(f"Pendiente de cobro: {pendiente:.2f} EUR"),
                       new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    pdf.set_font("Helvetica", "I", 8)
    pdf.multi_cell(0, 4, _latin1(
        "Este documento es un presupuesto/nota de honorarios. "
        "No constituye factura fiscal."
    ), new_x="LMARGIN", new_y="NEXT")
    return bytes(pdf.output())
